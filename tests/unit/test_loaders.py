"""
Unit tests for document loaders.
"""

import pytest
from pathlib import Path
import tempfile

import fitz

from app.rag.loaders import  TextLoader, MarkdownLoader, CSVLoader, DocxLoader, PDFLoader


class TestTextLoader:
    """Tests for TextLoader."""
    
    def test_load_simple_file(self):
        """Should load a simple text file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("This is a test document.")
            f.flush()
            
            loader = TextLoader()
            docs = loader.load(Path(f.name))
        
        assert len(docs) == 1
        assert docs[0].content == "This is a test document."
        assert docs[0].chunk_index == 0
    
    def test_chunking(self):
        """Should chunk large documents."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            # Write a long document
            f.write("This is sentence one. " * 100)
            f.flush()
            
            loader = TextLoader(chunk_size=100, chunk_overlap=20)
            docs = loader.load(Path(f.name))
        
        assert len(docs) > 1
        for doc in docs:
            assert len(doc.content) <= 120  # Some flexibility for word boundaries
    
    def test_file_not_found(self):
        """Should raise for missing file."""
        loader = TextLoader()
        
        with pytest.raises(FileNotFoundError):
            loader.load(Path("nonexistent.txt"))


class TestMarkdownLoader:
    """Tests for MarkdownLoader."""
    
    def test_split_by_headers(self):
        """Should split by markdown headers."""
        content = """# Header 1
Content under header 1.

## Header 2
Content under header 2.

## Header 3
Content under header 3.
"""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write(content)
            f.flush()
            
            loader = MarkdownLoader()
            docs = loader.load(Path(f.name))
        
        assert len(docs) == 3
        assert "Header 1" in docs[0].content
        assert "Header 2" in docs[1].content
        assert docs[1].metadata["section"] == "## Header 2"
    
    def test_no_headers(self):
        """Should handle markdown without headers."""
        content = "Just plain text without headers."
        
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write(content)
            f.flush()
            
            loader = MarkdownLoader()
            docs = loader.load(Path(f.name))
        
        assert len(docs) == 1
        assert docs[0].metadata["section"] == "intro"


class TestCSVLoader:
    """Tests for CSVLoader."""
    
    def test_load_csv(self):
        """Should load CSV with one doc per row."""
        content = """food,ph,aw
chicken,6.0,0.99
beef,5.5,0.98
"""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
            f.write(content)
            f.flush()
            
            loader = CSVLoader()
            docs = loader.load(Path(f.name))
        
        assert len(docs) == 2
        assert "chicken" in docs[0].content
        assert "6.0" in docs[0].content
    
    def test_content_columns(self):
        """Should use specified content columns."""
        content = """food,ph,aw,notes
chicken,6.0,0.99,raw
beef,5.5,0.98,ground
"""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
            f.write(content)
            f.flush()
            
            loader = CSVLoader(content_columns=["food", "notes"])
            docs = loader.load(Path(f.name))
        
        assert "food: chicken" in docs[0].content
        assert "notes: raw" in docs[0].content
        assert "ph" not in docs[0].content
    
    def test_metadata_columns(self):
        """Should include specified metadata columns."""
        content = """food,ph,aw
chicken,6.0,0.99
"""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
            f.write(content)
            f.flush()
            
            loader = CSVLoader(
                content_columns=["food"],
                metadata_columns=["ph", "aw"],
            )
            docs = loader.load(Path(f.name))
        
        assert docs[0].metadata["ph"] == "6.0"
        assert docs[0].metadata["aw"] == "0.99"
    
    def test_semicolon_delimiter(self):
        """Should handle different delimiters."""
        content = """food;ph;aw
chicken;6.0;0.99
"""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
            f.write(content)
            f.flush()
            
            loader = CSVLoader(delimiter=";")
            docs = loader.load(Path(f.name))
        
        assert len(docs) == 1
        assert "chicken" in docs[0].content


class TestChunking:
    """Tests for text chunking."""
    
    def test_chunk_at_sentence(self):
        """Should prefer breaking at sentence boundaries."""
        loader = TextLoader(chunk_size=50, chunk_overlap=10)
        
        text = "This is sentence one. This is sentence two. This is sentence three."
        chunks = loader.chunk_text(text, chunk_size=50, chunk_overlap=10)
        
        # Should break at period
        assert chunks[0].endswith(".")
    
    def test_small_text_no_chunking(self):
        """Should not chunk text smaller than chunk_size."""
        loader = TextLoader(chunk_size=100, chunk_overlap=10)
        
        text = "Short text."
        chunks = loader.chunk_text(text, chunk_size=100, chunk_overlap=10)
        
        assert len(chunks) == 1
        assert chunks[0] == "Short text."

class TestDocxLoader:
    """Tests for DocxLoader."""
    
    def test_load_docx(self):
        """Should load a DOCX file."""
        # Create a simple DOCX
        from docx import Document as DocxDocument
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = DocxDocument()
            doc.add_paragraph("This is paragraph one.")
            doc.add_paragraph("This is paragraph two.")
            doc.save(f.name)
            
            loader = DocxLoader()
            docs = loader.load(Path(f.name))
        
        assert len(docs) >= 1
        assert "paragraph one" in docs[0].content
        assert docs[0].metadata["format"] == "docx"
    
    def test_load_docx_with_table(self):
        """Should extract table content."""
        from docx import Document as DocxDocument
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = DocxDocument()
            doc.add_paragraph("Introduction text.")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Food"
            table.cell(0, 1).text = "pH"
            table.cell(1, 0).text = "Chicken"
            table.cell(1, 1).text = "6.0"
            doc.save(f.name)
            
            loader = DocxLoader(include_tables=True)
            docs = loader.load(Path(f.name))
        
        content = " ".join(d.content for d in docs)
        assert "Chicken" in content
        assert "6.0" in content
    
    def test_file_not_found(self):
        """Should raise for missing file."""
        loader = DocxLoader()
        
        with pytest.raises(FileNotFoundError):
            loader.load(Path("nonexistent.docx"))


class TestPDFLoader:
    """Tests for PDFLoader."""
    
    def test_load_pdf(self):
        """Should load a PDF file."""
        # Create temp file path, then close it so fitz can write
        temp_dir = tempfile.mkdtemp()
        pdf_path = Path(temp_dir) / "test.pdf"
        
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "This is test content in a PDF.")
        doc.save(str(pdf_path))
        doc.close()
        
        loader = PDFLoader()
        docs = loader.load(pdf_path)
        
        assert len(docs) >= 1
        assert "test content" in docs[0].content.lower()
        assert docs[0].metadata["format"] == "pdf"
        
        # Cleanup
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    def test_load_by_page(self):
        """Should load with page metadata."""
        temp_dir = tempfile.mkdtemp()
        pdf_path = Path(temp_dir) / "test.pdf"
        
        doc = fitz.open()
        page1 = doc.new_page()
        page1.insert_text((50, 50), "Page one content.")
        page2 = doc.new_page()
        page2.insert_text((50, 50), "Page two content.")
        doc.save(str(pdf_path))
        doc.close()
        
        loader = PDFLoader()
        docs = loader.load_by_page(pdf_path)
        
        assert len(docs) == 2
        assert docs[0].metadata["page"] == 1
        assert docs[1].metadata["page"] == 2
        
        # Cleanup
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    def test_file_not_found(self):
        """Should raise for missing file."""
        loader = PDFLoader()
        
        with pytest.raises(FileNotFoundError):
            loader.load(Path("nonexistent.pdf"))