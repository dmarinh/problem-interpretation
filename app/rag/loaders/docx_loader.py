"""
DOCX loader using python-docx.
"""

from pathlib import Path

from docx import Document as DocxDocument

from app.config import settings
from app.rag.loaders.base import BaseLoader, Document


class DocxLoader(BaseLoader):
    """
    Loader for Microsoft Word documents (.docx).
    
    Extracts text from paragraphs and tables.
    """
    
    def __init__(
        self,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
        include_tables: bool = True,
    ):
        """
        Initialize DOCX loader.
        
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between chunks
            include_tables: Whether to extract text from tables
        """
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
        self.include_tables = include_tables
    
    def load(self, file_path: Path) -> list[Document]:
        """Load and chunk a DOCX file."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        doc = DocxDocument(file_path)
        
        # Extract paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract tables
        if self.include_tables:
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    text_parts.append(table_text)
        
        text = "\n\n".join(text_parts)
        
        if not text.strip():
            return []
        
        chunks = self.chunk_text(text, self.chunk_size, self.chunk_overlap)
        
        documents = []
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                content=chunk,
                metadata={"filename": file_path.name, "format": "docx"},
                source=str(file_path),
                chunk_index=i,
            ))
        
        return documents
    
    def _extract_table(self, table) -> str:
        """Extract text from a table as formatted rows."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        return "\n".join(rows)